# file discovery, type classification, and corpus health checks
# tentaqles.graph.native — ported from graphify with all patches as first-class code
from __future__ import annotations
import fnmatch
import json
import os
import re
from enum import Enum
from pathlib import Path

import pathspec


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"


_MANIFEST_PATH = "tentaqles-out/manifest.json"

CODE_EXTENSIONS = {'.py', '.ts', '.js', '.jsx', '.tsx', '.go', '.rs', '.java', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp', '.rb', '.swift', '.kt', '.kts', '.cs', '.scala', '.php', '.lua', '.toc', '.zig', '.ps1', '.ex', '.exs', '.m', '.mm', '.jl'}
DOC_EXTENSIONS = {'.md', '.txt', '.rst'}
PAPER_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}
OFFICE_EXTENSIONS = {'.docx', '.xlsx', '.pptx'}

CORPUS_WARN_THRESHOLD = 50_000    # words - below this, warn "you may not need a graph"
CORPUS_UPPER_THRESHOLD = 500_000  # words - above this, warn about token cost
FILE_COUNT_UPPER = 200             # files - above this, warn about token cost

# Files that may contain secrets - skip silently
_SENSITIVE_PATTERNS = [
    re.compile(r'(^|[\\/])\.(env|envrc)(\.|$)', re.IGNORECASE),
    re.compile(r'\.(pem|key|p12|pfx|cert|crt|der|p8)$', re.IGNORECASE),
    re.compile(r'(credential|secret|passwd|password|token|private_key)', re.IGNORECASE),
    re.compile(r'(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$'),
    re.compile(r'(\.netrc|\.pgpass|\.htpasswd)$', re.IGNORECASE),
    re.compile(r'(aws_credentials|gcloud_credentials|service.account)', re.IGNORECASE),
]

# Signals that a .md/.txt file is actually a converted academic paper
_PAPER_SIGNALS = [
    re.compile(r'\barxiv\b', re.IGNORECASE),
    re.compile(r'\bdoi\s*:', re.IGNORECASE),
    re.compile(r'\babstract\b', re.IGNORECASE),
    re.compile(r'\bproceedings\b', re.IGNORECASE),
    re.compile(r'\bjournal\b', re.IGNORECASE),
    re.compile(r'\bpreprint\b', re.IGNORECASE),
    re.compile(r'\\cite\{'),          # LaTeX citation
    re.compile(r'\[\d+\]'),           # Numbered citation [1], [23] (inline)
    re.compile(r'\[\n\d+\n\]'),       # Numbered citation spread across lines (markdown conversion)
    re.compile(r'eq\.\s*\d+|equation\s+\d+', re.IGNORECASE),
    re.compile(r'\d{4}\.\d{4,5}'),   # arXiv ID like 1706.03762
    re.compile(r'\bwe propose\b', re.IGNORECASE),   # common academic phrasing
    re.compile(r'\bliterature\b', re.IGNORECASE),   # "from the literature"
]
_PAPER_SIGNAL_THRESHOLD = 3  # need at least this many signals to call it a paper


def _is_sensitive(path: Path) -> bool:
    """Return True if this file likely contains secrets and should be skipped."""
    name = path.name
    full = str(path)
    return any(p.search(name) or p.search(full) for p in _SENSITIVE_PATTERNS)


def _looks_like_paper(path: Path) -> bool:
    """Heuristic: does this text file read like an academic paper?"""
    try:
        # Only scan first 3000 chars for speed
        text = path.read_text(errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}


def classify_file(path: Path) -> FileType | None:
    ext = path.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        # PDFs inside Xcode asset catalogs are vector icons, not papers
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DOC_EXTENSIONS:
        # Check if it's a converted paper
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Skip entirely empty rows
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def convert_with_docling(source: Path, output_dir: Path) -> Path | None:
    """Convert a document to markdown using docling.

    Supports PDF, DOCX, XLSX, PPTX and other formats docling handles.
    Returns the path to the saved .md file, or None if docling is not
    installed or conversion fails.
    """
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return None

    try:
        ext = source.suffix.lower()

        if ext == ".pdf":
            # Enable image extraction for PDFs
            from docling.document_converter import PdfFormatOption
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions

            pipeline_options = PdfPipelineOptions()
            pipeline_options.generate_picture_images = True
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
        else:
            converter = DocumentConverter()

        result = converter.convert(str(source))
        md_text = result.document.export_to_markdown()

        if not md_text.strip():
            return None

        output_dir.mkdir(parents=True, exist_ok=True)
        out_path = output_dir / (source.stem + ".md")
        # Avoid collisions by adding hash if file already exists with different source
        if out_path.exists():
            import hashlib
            name_hash = hashlib.sha256(str(source.resolve()).encode()).hexdigest()[:8]
            out_path = output_dir / f"{source.stem}_{name_hash}.md"

        out_path.write_text(
            f"<!-- converted from {source.name} via docling -->\n\n{md_text}",
            encoding="utf-8",
        )
        return out_path
    except Exception:
        return None


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions
    import hashlib
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        if ext == ".pptx":
            # Try docling for PPTX word count; fall back to 0
            try:
                from docling.document_converter import DocumentConverter
                result = DocumentConverter().convert(str(path))
                return len(result.document.export_to_markdown().split())
            except Exception:
                return 0
        return len(path.read_text(errors="ignore").split())
    except Exception:
        return 0


# Directory names to always skip - venvs, caches, build artifacts, deps
_SKIP_DIRS = {
    "venv", ".venv", "env", ".env",
    "node_modules", "__pycache__", ".git",
    "dist", "build", "target", "out",
    "site-packages", "lib64",
    ".pytest_cache", ".mypy_cache", ".ruff_cache",
    ".tox", ".eggs", "*.egg-info",
}

def _is_noise_dir(part: str) -> bool:
    """Return True if this directory name looks like a venv, cache, or dep dir."""
    if part in _SKIP_DIRS:
        return True
    # Catch *_venv, *_repo/site-packages patterns
    if part.endswith("_venv") or part.endswith("_env"):
        return True
    if part.endswith(".egg-info"):
        return True
    return False


def _load_ignore_patterns(directory: Path) -> list[str]:
    """Read .gitignore and .tentaqlesignore from a directory and return combined patterns.

    Both files use gitignore semantics. .tentaqlesignore patterns are appended
    after .gitignore patterns so they take precedence (last match wins in
    pathspec, and negation patterns can override earlier ones).
    """
    patterns: list[str] = []
    for filename in (".gitignore", ".tentaqlesignore"):
        ignore_file = directory / filename
        if not ignore_file.exists():
            continue
        for line in ignore_file.read_text(errors="ignore").splitlines():
            line = line.strip()
            if line and not line.startswith("#"):
                patterns.append(line)
    return patterns


class _IgnoreTree:
    """Cascading ignore matcher that mirrors git's nested .gitignore behavior.

    At each directory level, patterns from .gitignore and .tentaqlesignore are
    collected. To check a path, all ancestor patterns from root down to the
    path's parent are combined into a single pathspec — this lets negation
    patterns in child directories correctly override parent ignores.
    """

    def __init__(self, root: Path) -> None:
        self.root = root
        # Cache: directory path -> raw pattern lines from that dir's ignore files
        self._raw_patterns: dict[Path, list[str]] = {}
        # Cache: directory path -> combined GitIgnoreSpec (root..dir merged)
        self._merged_specs: dict[Path, pathspec.GitIgnoreSpec | None] = {}
        self._pattern_count = 0
        # Pre-load root level
        self._ensure_loaded(root)

    def _ensure_loaded(self, directory: Path) -> None:
        """Load ignore patterns for a directory if not already cached."""
        if directory in self._raw_patterns:
            return
        patterns = _load_ignore_patterns(directory)
        self._raw_patterns[directory] = patterns
        self._pattern_count += len(patterns)

    def _get_merged_spec(self, directory: Path) -> pathspec.GitIgnoreSpec | None:
        """Get a merged spec combining all patterns from root down to directory.

        Patterns from each level are rewritten as relative to root before merging,
        so the final spec can match paths relative to root. Parent patterns come
        first, child patterns last — pathspec uses last-match-wins, so child
        negation patterns correctly override parent ignores.
        """
        if directory in self._merged_specs:
            return self._merged_specs[directory]

        # Build the chain: root -> ... -> directory
        try:
            rel_to_root = directory.relative_to(self.root)
        except ValueError:
            self._merged_specs[directory] = None
            return None

        chain = [self.root]
        current = self.root
        for part in rel_to_root.parts:
            current = current / part
            chain.append(current)

        # Collect all patterns, rewriting each level's patterns to be root-relative
        all_patterns: list[str] = []
        for ancestor in chain:
            raw = self._raw_patterns.get(ancestor, [])
            if not raw:
                continue
            if ancestor == self.root:
                all_patterns.extend(raw)
            else:
                prefix = str(ancestor.relative_to(self.root)).replace(os.sep, "/")
                for pat in raw:
                    negated = pat.startswith("!")
                    clean = pat[1:] if negated else pat
                    # Rewrite pattern to be root-relative
                    rewritten = f"{prefix}/{clean}"
                    if negated:
                        rewritten = f"!{rewritten}"
                    all_patterns.append(rewritten)

        if all_patterns:
            spec = pathspec.GitIgnoreSpec.from_lines(all_patterns)
        else:
            spec = None

        self._merged_specs[directory] = spec
        return spec

    def notify_enter_dir(self, directory: Path) -> None:
        """Called when os.walk enters a directory. Loads its ignore files."""
        self._ensure_loaded(directory)

    @property
    def pattern_count(self) -> int:
        return self._pattern_count

    def is_ignored(self, path: Path, is_dir: bool = False) -> bool:
        """Check if path is ignored by the combined ancestor ignore patterns."""
        try:
            rel = path.relative_to(self.root)
        except ValueError:
            return False

        rel_str = str(rel).replace(os.sep, "/")
        if is_dir:
            rel_str += "/"

        spec = self._get_merged_spec(path.parent)
        if spec is None:
            return False

        return spec.match_file(rel_str)


# Backwards-compatible wrappers used by _is_ignored calls that still pass
# a flat pattern list (e.g. external callers). These are only used as fallback.
def _load_tentaqlesignore(root: Path) -> list[str]:
    """Legacy: read .tentaqlesignore from root. Kept for API compat."""
    return _load_ignore_patterns(root)


def _is_ignored(path: Path, root: Path, patterns: list[str]) -> bool:
    """Legacy fallback: match against a flat pattern list using fnmatch."""
    if not patterns:
        return False
    try:
        rel = str(path.relative_to(root))
    except ValueError:
        return False
    rel = rel.replace(os.sep, "/")
    parts = rel.split("/")
    for pattern in patterns:
        p = pattern.strip("/")
        if not p:
            continue
        if fnmatch.fnmatch(rel, p):
            return True
        if fnmatch.fnmatch(path.name, p):
            return True
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, p):
                return True
            if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                return True
    return False


def detect(root: Path, *, follow_symlinks: bool = False) -> dict:
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
    }
    total_words = 0

    skipped_sensitive: list[str] = []
    ignore_tree = _IgnoreTree(root)

    # Always include tentaqles-out/memory/ - query results filed back into the graph
    memory_dir = root / "tentaqles-out" / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []

    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(scan_root, followlinks=follow_symlinks):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                # Load ignore files for this directory level
                ignore_tree.notify_enter_dir(dp)
                # Prune noise dirs in-place so os.walk never descends into them
                dirnames[:] = [
                    d for d in dirnames
                    if not d.startswith(".")
                    and not _is_noise_dir(d)
                    and not ignore_tree.is_ignored(dp / d, is_dir=True)
                ]
            for fname in filenames:
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    converted_dir = root / "tentaqles-out" / "converted"

    for p in all_files:
        # For memory dir files, skip hidden/noise filtering
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            # Hidden files are already excluded via dir pruning above,
            # but catch hidden files at the root level
            if p.name.startswith("."):
                continue
            # Skip files inside our own converted/ dir (avoid re-processing sidecars)
            if str(p).startswith(str(converted_dir)):
                continue
        if ignore_tree.is_ignored(p):
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            continue
        ftype = classify_file(p)
        if ftype:
            # Office files: convert to markdown sidecar so subagents can read them
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                ext_lower = p.suffix.lower()
                md_path = None
                if ext_lower == ".pptx":
                    # PPTX: docling first (python-docx doesn't handle pptx)
                    md_path = convert_with_docling(p, converted_dir)
                    if not md_path:
                        md_path = convert_office_file(p, converted_dir)
                else:
                    # DOCX/XLSX: existing converter first, docling as fallback
                    md_path = convert_office_file(p, converted_dir)
                    if not md_path:
                        md_path = convert_with_docling(p, converted_dir)
                if md_path:
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    # Conversion failed (library not installed) - skip with note
                    skipped_sensitive.append(str(p) + " [office conversion failed - pip install tentaqles[office]]")
                continue

            # PDFs: try docling for richer markdown extraction
            if ftype == FileType.PAPER and p.suffix.lower() == ".pdf":
                md_path = convert_with_docling(p, converted_dir)
                if md_path:
                    # Use the richer markdown version instead of raw PDF
                    files[FileType.DOCUMENT].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    # Docling not available or failed - fall back to raw PDF
                    files[ftype].append(str(p))
                    total_words += count_words(p)
                continue

            files[ftype].append(str(p))
            total_words += count_words(p)

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD

    # Determine warning - lower bound, upper bound, or sensitive files skipped
    warning: str | None = None
    if not needs_graph:
        warning = (
            f"Corpus is ~{total_words:,} words - fits in a single context window. "
            f"You may not need a graph."
        )
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive (many Claude tokens). "
            f"Consider running on a subfolder, or use --no-semantic to run AST-only."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "ignore_patterns": ignore_tree.pattern_count,
    }


def load_manifest(manifest_path: str = _MANIFEST_PATH) -> dict[str, float]:
    """Load the file modification time manifest from a previous run."""
    try:
        return json.loads(Path(manifest_path).read_text())
    except Exception:
        return {}


def save_manifest(files: dict[str, list[str]], manifest_path: str = _MANIFEST_PATH) -> None:
    """Save current file mtimes so the next --update run can diff against them."""
    manifest: dict[str, float] = {}
    for file_list in files.values():
        for f in file_list:
            try:
                manifest[f] = Path(f).stat().st_mtime
            except OSError:
                pass  # file deleted between detect() and manifest write - skip it
    Path(manifest_path).parent.mkdir(parents=True, exist_ok=True)
    Path(manifest_path).write_text(json.dumps(manifest, indent=2))


def detect_incremental(root: Path, manifest_path: str = _MANIFEST_PATH) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    Compares current file mtimes against the stored manifest.
    Use for --update mode: re-extract only what changed, merge into existing graph.
    """
    full = detect(root)
    manifest = load_manifest(manifest_path)

    if not manifest:
        # No previous run - treat everything as new
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored_mtime = manifest.get(f)
            try:
                current_mtime = Path(f).stat().st_mtime
            except Exception:
                current_mtime = 0
            if stored_mtime is None or current_mtime > stored_mtime:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    # Files in manifest that no longer exist - their cached nodes are now ghost nodes
    current_files = {f for flist in full["files"].values() for f in flist}
    deleted_files = [f for f in manifest if f not in current_files]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
